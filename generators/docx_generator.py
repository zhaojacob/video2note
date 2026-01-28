"""
Word document generator - New format with proper fonts
字体规范：
- 英文字体：Times New Roman
- 中文字体：宋体
- 字号：小四 (12pt)
- 颜色：黑色正文，灰色图片标注
"""
import re
from pathlib import Path
from typing import Dict, Any, Optional, List

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

from utils.logger import get_logger
from utils.file_handler import ensure_dir

logger = get_logger(__name__)

# Font constants
FONT_CHINESE = "宋体"
FONT_ENGLISH = "Times New Roman"
FONT_SIZE_NORMAL = Pt(12)  # 小四
FONT_SIZE_TITLE = Pt(18)
FONT_SIZE_HEADING = Pt(14)
FONT_SIZE_CAPTION = Pt(10)
COLOR_BLACK = RGBColor(0, 0, 0)
COLOR_GRAY = RGBColor(128, 128, 128)


class DocxGenerator:
    """Generate Word documents from structured data with proper fonts"""

    def __init__(self, output_dir: Optional[Path] = None):
        """
        Initialize DOCX generator

        Args:
            output_dir: Output directory for documents
        """
        self.output_dir = output_dir or Path("output/notes")
        ensure_dir(self.output_dir)

    def generate(
        self,
        data: Dict[str, Any],
        filename: str = None
    ) -> Path:
        """
        Generate Word document with new structure

        Args:
            data: Structured note data
            filename: Output filename (without .docx)

        Returns:
            Path to generated document
        """
        logger.info("Generating Word document (new format)")

        # Create document
        doc = Document()

        # Set document properties
        self._setup_styles(doc)

        # Add title (with translation if available)
        self._add_title(
            doc, 
            data["metadata"]["title"],
            data["metadata"].get("title_translated", "")
        )

        # Add metadata line (时间 | 来源 | 作者 | 链接)
        self._add_metadata_line(doc, data["metadata"])

        # Add version line
        self._add_version_line(doc, data["metadata"])

        # Add summary section (with translation if available)
        self._add_summary(
            doc, 
            data.get("summary", ""),
            data.get("summary_translated", "")
        )

        # Add transcript section with images
        # Use chapters if available (with titles), otherwise full_transcript with timestamps
        self._add_transcript_section(
            doc,
            data.get("full_transcript", []),
            data.get("sections", []),
            data.get("polished_text", ""),
            data.get("chapters", []),
            data.get("heading_markers", []),
            data.get("statistics", {}).get("total_frames", 0),
            data.get("structured_sections", [])
        )

        # Add statistics
        self._add_statistics(doc, data["statistics"])

        # Save document
        if filename is None:
            filename = self._sanitize_filename(data["metadata"]["title"])

        # Add timestamp suffix (YYYYMMDD_HHMMSS)
        from datetime import datetime
        timestamp_suffix = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = self.output_dir / f"{filename}_{timestamp_suffix}.docx"
        doc.save(str(output_path))

        logger.info(f"Word document saved: {output_path}")

        return output_path

    def _setup_styles(self, doc: Document):
        """Setup document styles with proper fonts"""
        # Default font - 小四 (12pt)
        style = doc.styles['Normal']
        font = style.font
        font.name = FONT_ENGLISH  # Times New Roman for English
        font.size = FONT_SIZE_NORMAL
        font.color.rgb = COLOR_BLACK

        # Set Chinese font - 宋体
        style.element.rPr.rFonts.set(
            qn('w:eastAsia'),
            FONT_CHINESE
        )

    def _set_run_font(self, run, size=None, color=None, bold=False, italic=False):
        """Helper to set run font properties"""
        run.font.name = FONT_ENGLISH
        run.font.size = size or FONT_SIZE_NORMAL
        run.font.color.rgb = color or COLOR_BLACK
        run.font.bold = bold
        run.font.italic = italic
        # Set Chinese font
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CHINESE)

    def _add_title(self, doc: Document, title: str, title_translated: str = ""):
        """Add document title with optional translation"""
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for run in heading.runs:
            self._set_run_font(run, size=FONT_SIZE_TITLE, bold=True)
        
        # Add translated title if available
        if title_translated:
            p = doc.add_paragraph()
            run = p.add_run(title_translated)
            self._set_run_font(run, size=Pt(14), color=COLOR_GRAY, italic=True)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def _add_metadata_line(self, doc: Document, metadata: Dict[str, Any]):
        """Add single metadata line: 时间 | 来源 | 作者 | 链接"""
        p = doc.add_paragraph()
        
        parts = []
        
        # 时间
        created_at = metadata.get('created_at', '')
        if created_at:
            date_str = created_at[:10] if len(created_at) >= 10 else created_at
            parts.append(date_str)
        
        # 来源
        source = metadata.get('source', 'unknown')
        parts.append(source.capitalize())
        
        # 作者
        uploader = metadata.get('uploader', '')
        if uploader:
            parts.append(uploader)
        
        # 链接
        url = metadata.get('url', '')
        if url:
            parts.append(url)
        
        metadata_text = " | ".join(parts)
        run = p.add_run(metadata_text)
        self._set_run_font(run, size=FONT_SIZE_CAPTION, color=COLOR_GRAY)
        
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def _add_version_line(self, doc: Document, metadata: Dict[str, Any]):
        """Add version line"""
        p = doc.add_paragraph()
        
        version = metadata.get('version', 'v1.0')
        created_at = metadata.get('created_at', '')
        if created_at:
            created_at = created_at[:19].replace('T', ' ')
        
        run = p.add_run(f"生成版本: {version} - {created_at}")
        self._set_run_font(run, size=FONT_SIZE_CAPTION, color=COLOR_GRAY, italic=True)
        
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()  # Empty line

    def _clean_markdown_formatting(self, text: str) -> str:
        """
        Clean markdown formatting from summary text for Word document

        Args:
            text: Raw text with potential markdown formatting

        Returns:
            Cleaned text suitable for Word document
        """
        if not text:
            return text

        # Remove markdown bold markers
        text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)  # ***bold***
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)      # **bold**
        text = re.sub(r'\*(.+?)\*', r'\1', text)          # *italic*

        # Clean up label formatting like "*   **Title:** text"
        text = re.sub(r'^\*\s+\*\*([^:]+):\*\*\s*', r'\1：', text, flags=re.MULTILINE)

        # Remove bullet point markers if they exist (Word uses its own formatting)
        text = re.sub(r'^\s*[\*\-•]\s+', '• ', text, flags=re.MULTILINE)

        # Clean up extra spaces
        text = re.sub(r'\s{3,}', ' ', text)  # Replace 3+ spaces with single space
        text = re.sub(r'\n\s+\n', '\n\n', text)  # Remove trailing spaces from empty lines

        return text.strip()

    def _add_summary(self, doc: Document, summary: str, summary_translated: str = ""):
        """Add summary section with optional translation"""
        heading = doc.add_heading("摘要", level=2)
        for run in heading.runs:
            self._set_run_font(run, size=FONT_SIZE_HEADING, bold=True)

        if summary:
            # Clean markdown formatting before adding to document
            cleaned_summary = self._clean_markdown_formatting(summary)

            p = doc.add_paragraph()
            run = p.add_run(cleaned_summary)
            self._set_run_font(run)
            p.paragraph_format.line_spacing = 1.5

            # Add translated summary if available
            if summary_translated:
                cleaned_translated = self._clean_markdown_formatting(summary_translated)
                p2 = doc.add_paragraph()
                run2 = p2.add_run(cleaned_translated)
                self._set_run_font(run2, color=COLOR_GRAY, italic=True)
                p2.paragraph_format.line_spacing = 1.5
        else:
            p = doc.add_paragraph()
            run = p.add_run("（摘要生成中或未配置 DeepSeek API）")
            self._set_run_font(run, color=COLOR_GRAY, italic=True)

        doc.add_paragraph()

    def _add_transcript_section(
        self,
        doc: Document,
        full_transcript: List[Dict[str, Any]],
        sections: List[Dict[str, Any]],
        polished_text: str = "",
        chapters: List[Dict[str, Any]] = None,
        heading_markers: List[Dict[str, Any]] = None,
        total_frames: int = 0,
        structured_sections: List[Dict[str, Any]] = None
    ):
        """Add transcript section with images"""
        heading = doc.add_heading("正文", level=2)
        for run in heading.runs:
            self._set_run_font(run, size=FONT_SIZE_HEADING, bold=True)

        # Extract images from full_transcript for later insertion
        images = []
        if full_transcript:
            images = [item for item in full_transcript if item.get("type") == "image"]

        # Priority 0: Use structured sections (JSON pipeline)
        if structured_sections:
            self._render_structured_sections(doc, structured_sections)
            return

        # Prefer chapters (with titles and structured content)
        if chapters:
            self._render_chapters_with_images(doc, chapters, images)
        # NEW: Use heading markers with full_transcript (has timestamps + images aligned)
        elif heading_markers and full_transcript:
            self._render_text_with_markers_aligned(doc, full_transcript, heading_markers)
        # Fallback: heading markers with polished text (no timestamps)
        elif heading_markers and polished_text:
            self._render_polished_text_with_images(doc, polished_text, images, heading_markers)
        # Fallback to polished text (with punctuation, paragraphs)
        elif polished_text:
            self._render_polished_text_with_images(doc, polished_text, images)
        # Fallback to full_transcript format (has timestamps and images)
        elif full_transcript:
            self._render_full_transcript(doc, full_transcript)
        elif sections:
            # Fallback to old sections format
            self._render_sections_as_transcript(doc, sections)
        else:
            p = doc.add_paragraph()
            run = p.add_run("（无转录内容）")
            self._set_run_font(run, color=COLOR_GRAY, italic=True)

    def _render_structured_sections(self, doc: Document, sections: List[Dict[str, Any]]):
        """Render structured sections with headers, timestamps and images"""
        for section in sections:
            # Section Header
            title = section.get("title", "")
            if title:
                heading = doc.add_heading(title, level=3)
                self._style_heading(heading)
            
            for para in section.get("paragraphs", []):
                # Images (render BEFORE text to ensure opening images appear at top)
                images = para.get("images", [])
                for img in images:
                    self._add_transcript_image(doc, img)

                # Paragraph
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Inches(0.5)
                p.paragraph_format.line_spacing = 1.5
                
                # Timestamp
                timestamp = para.get("timestamp", "")
                if timestamp:
                    ts_run = p.add_run(f"[{timestamp}] ")
                    self._set_run_font(ts_run, bold=True)
                
                # Content
                content = para.get("content", "").strip()
                if content:
                    # 替换文本中的 \n 为空格，避免软换行符 ^l
                    content_clean = content.replace('\n', ' ').replace('\r', '')
                    run = p.add_run(content_clean)
                    self._set_run_font(run)
                
                # Translation (if available) - render as separate paragraph after original
                content_translated = para.get("content_translated", "").strip()
                if content_translated:
                    content_translated_clean = content_translated.replace('\n', ' ').replace('\r', '')
                    p2 = doc.add_paragraph()
                    p2.paragraph_format.left_indent = Inches(0.3)
                    p2.paragraph_format.line_spacing = 1.5
                    trans_run = p2.add_run(content_translated_clean)
                    self._set_run_font(trans_run, color=COLOR_GRAY, italic=True)

    def _render_chapters_with_images(
        self, 
        doc: Document, 
        chapters: List[Dict[str, Any]],
        images: List[Dict[str, Any]]
    ):
        """Render chapters with titles, content, and distributed images"""
        from docx.shared import RGBColor
        
        num_chapters = len(chapters)
        num_images = len(images)
        
        # Distribute images evenly across chapters
        images_per_chapter = max(1, num_images // num_chapters) if num_chapters > 0 else num_images
        image_index = 0
        
        for i, chapter in enumerate(chapters):
            title = chapter.get('title', f'章节 {i + 1}')
            content = chapter.get('content', '')
            
            # Add chapter title as heading level 3
            chapter_heading = doc.add_heading(title, level=3)
            self._style_heading(chapter_heading)
            
            # Add images for this chapter (distribute evenly)
            chapter_images_count = images_per_chapter
            # Last chapter gets remaining images
            if i == num_chapters - 1:
                chapter_images_count = num_images - image_index
            
            for j in range(chapter_images_count):
                if image_index < num_images:
                    self._add_transcript_image(doc, images[image_index])
                    image_index += 1
            
            # Add chapter content
            if content:
                paragraphs = content.split('\n\n') if '\n\n' in content else content.split('\n')
                
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if not para_text:
                        continue
                    
                    # 替换文本中的 \n 为空格，避免软换行符
                    para_text_clean = para_text.replace('\n', ' ')
                    p = doc.add_paragraph()
                    p.paragraph_format.first_line_indent = Inches(0.5)
                    p.paragraph_format.line_spacing = 1.5
                    
                    run = p.add_run(para_text_clean)
                    self._set_run_font(run)
            
            doc.add_paragraph()

    def _render_polished_text_with_images(
        self,
        doc: Document,
        text: str,
        images: List[Dict[str, Any]],
        heading_markers: List[Dict[str, Any]] = None
    ):
        """Render polished text with images inserted"""
        # Priority 1: Use heading markers if available (NEW!)
        if heading_markers:
            self._render_text_with_markers(doc, text, images, heading_markers)
            return

        # Priority 2: Check if text contains chapter markers (legacy)
        if '## ' in text:
            from utils.text_polisher import TextPolisher
            polisher = TextPolisher()
            chapters = polisher.extract_chapters(text)
            if chapters:
                self._render_chapters_with_images(doc, chapters, images)
                return

        # Priority 3: Add all images at the beginning
        for image in images:
            self._add_transcript_image(doc, image)

        # Then add text paragraphs
        paragraphs = text.split('\n\n') if '\n\n' in text else text.split('\n')

        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            # 替换文本中的 \n 为空格，避免软换行符 ^l
            para_text_clean = para_text.replace('\n', ' ').replace('\r', '')
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.line_spacing = 1.5

            run = p.add_run(para_text_clean)
            self._set_run_font(run)

    def _render_text_with_markers(
        self,
        doc: Document,
        text: str,
        images: List[Dict[str, Any]],
        heading_markers: List[Dict[str, Any]]
    ):
        """
        Render text with heading markers inserted at paragraph positions

        Args:
            doc: Document object
            text: Full polished text (paragraph-separated by \\n\\n)
            images: List of image items to distribute
            heading_markers: List of {"title": str, "paragraph_index": int}
        """
        paragraphs = text.split('\n\n') if '\n\n' in text else text.split('\n')

        # Distribute images across heading markers
        images_per_marker = max(1, len(images) // len(heading_markers)) if heading_markers else len(images)
        image_idx = 0

        for i, para_text in enumerate(paragraphs):
            # Check if we need to insert a heading at this paragraph
            for marker in heading_markers:
                if marker["paragraph_index"] == i:
                    # Insert heading
                    heading = doc.add_heading(marker["title"], level=3)
                    self._style_heading(heading)

                    # Add images for this section
                    for _ in range(images_per_marker):
                        if image_idx < len(images):
                            self._add_transcript_image(doc, images[image_idx])
                            image_idx += 1

            # Add paragraph
            if para_text.strip():
                # 替换文本中的 \n 为空格，避免软换行符
                para_text_clean = para_text.strip().replace('\n', ' ')
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Inches(0.5)
                p.paragraph_format.line_spacing = 1.5
                run = p.add_run(para_text_clean)
                self._set_run_font(run)

    def _render_text_with_markers_aligned(
        self,
        doc: Document,
        full_transcript: List[Dict[str, Any]],
        heading_markers: List[Dict[str, Any]]
    ):
        """
        Render full_transcript with heading markers inserted at paragraph positions

        This method uses full_transcript (which has timestamps) instead of polished_text,
        ensuring that:
        1. All paragraphs have timestamps at the beginning
        2. Images are aligned with paragraphs based on their timestamps

        Args:
            doc: Document object
            full_transcript: Items with timestamps (already sorted with images interleaved)
            heading_markers: List of {"title": str, "paragraph_index": int}
        """
        # Step 1: Group consecutive text items into paragraphs
        # This uses the same 3-second merging logic as the text processor
        paragraphs_with_timestamps = []
        current_para = None

        for item in full_transcript:
            if item["type"] == "text":
                # Check if we should merge with current paragraph
                # Merge consecutive items within 3 seconds
                if current_para:
                    time_gap = item["timestamp"] - current_para["end_time"]
                    if time_gap < 3.0:
                        # Merge with current paragraph
                        current_para["content"] += " " + item["content"]
                        current_para["end_time"] = item.get("end_time", item["timestamp"])
                        current_para["timestamp_formatted"] = item.get("timestamp_formatted", "")
                        continue

                # Save current paragraph if exists
                if current_para:
                    paragraphs_with_timestamps.append(current_para)

                # Start new paragraph
                current_para = {
                    "type": "text",
                    "timestamp": item["timestamp"],
                    "end_time": item.get("end_time", item["timestamp"]),
                    "content": item.get("content", ""),
                    "timestamp_formatted": item.get("timestamp_formatted", ""),
                    "content_translated": item.get("content_translated", "")
                }
            else:  # Image
                # Flush paragraph before image
                if current_para:
                    paragraphs_with_timestamps.append(current_para)
                    current_para = None
                # Add image as separate item
                paragraphs_with_timestamps.append(item)

        # Don't forget last paragraph
        if current_para:
            paragraphs_with_timestamps.append(current_para)

        # Step 2: Render with headings and proper image alignment
        para_idx = 0
        for item in paragraphs_with_timestamps:
            if item["type"] == "text":
                # Check if heading should be inserted before this paragraph
                for marker in heading_markers:
                    if marker["paragraph_index"] == para_idx:
                        # Insert heading
                        heading = doc.add_heading(marker["title"], level=3)
                        self._style_heading(heading)

                # Add paragraph WITH timestamp
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Inches(0.5)
                p.paragraph_format.line_spacing = 1.5

                # Add timestamp in bold
                timestamp = item.get("timestamp_formatted", "")
                if not timestamp:
                    timestamp = f"[{self._format_timestamp(item['timestamp'])}]"
                ts_run = p.add_run(f"{timestamp} ")
                self._set_run_font(ts_run, bold=True)

                # Add content
                content = item.get("content", "").strip()
                if content:
                    # 替换文本中的 \n 为空格，避免软换行符 ^l
                    content_clean = content.replace('\n', ' ').replace('\r', '')
                    content_run = p.add_run(content_clean)
                    self._set_run_font(content_run)

                # Add translation if available
                content_translated = item.get("content_translated", "").strip()
                if content_translated:
                    content_translated_clean = content_translated.replace('\n', ' ').replace('\r', '')
                    p2 = doc.add_paragraph()
                    p2.paragraph_format.left_indent = Inches(0.3)
                    p2.paragraph_format.line_spacing = 1.5
                    trans_run = p2.add_run(content_translated_clean)
                    self._set_run_font(trans_run, color=COLOR_GRAY, italic=True)

                para_idx += 1

            elif item["type"] == "image":
                # Add image (already at correct position in list based on timestamp)
                self._add_transcript_image(doc, item)

    def _style_heading(self, heading):
        """Apply consistent styling to level 3 headings"""
        for run in heading.runs:
            run.font.name = FONT_ENGLISH
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CHINESE)

    def _render_chapters(self, doc: Document, chapters: List[Dict[str, Any]]):
        """Render chapters with titles and content (no images)"""
        self._render_chapters_with_images(doc, chapters, [])

    def _render_polished_text(self, doc: Document, text: str):
        """Render polished transcript text with proper paragraphs (no images)"""
        self._render_polished_text_with_images(doc, text, [])

    def _render_full_transcript(self, doc: Document, items: List[Dict[str, Any]]):
        """Render full transcript with embedded images and translations"""
        for item in items:
            if item["type"] == "text":
                # 优先使用预格式化的时间戳
                timestamp = item.get("timestamp_formatted", "")
                if not timestamp:
                    timestamp = f"[{self._format_timestamp(item['timestamp'])}]"

                content = item.get("content", "").strip()
                content_translated = item.get("content_translated", "").strip()

                if content:
                    # 替换文本中的 \n 为空格，避免软换行符 ^l
                    content_clean = content.replace('\n', ' ').replace('\r', '')
                    
                    p = doc.add_paragraph()
                    # Timestamp in bold
                    ts_run = p.add_run(f"{timestamp} ")
                    self._set_run_font(ts_run, bold=True)
                    # Content
                    content_run = p.add_run(content_clean)
                    self._set_run_font(content_run)
                    p.paragraph_format.line_spacing = 1.5

                    # Add translation if available
                    if content_translated:
                        content_translated_clean = content_translated.replace('\n', ' ').replace('\r', '')
                        p2 = doc.add_paragraph()
                        # Indent for translation
                        p2.paragraph_format.left_indent = Inches(0.3)
                        trans_run = p2.add_run(content_translated_clean)
                        self._set_run_font(trans_run, color=COLOR_GRAY, italic=True)
                        p2.paragraph_format.line_spacing = 1.5

            elif item["type"] == "image":
                self._add_transcript_image(doc, item)

    def _render_sections_as_transcript(self, doc: Document, sections: List[Dict[str, Any]]):
        """Fallback: render old sections format as transcript"""
        for section in sections:
            timestamp = self._format_timestamp(section.get('start_time', 0))
            content = section.get('content', '').strip()
            
            # Add images first
            for image in section.get('images', []):
                self._add_image(doc, image, section.get('start_time', 0))
            
            # Add text content
            if content:
                p = doc.add_paragraph()
                ts_run = p.add_run(f"[{timestamp}] ")
                self._set_run_font(ts_run, bold=True)
                content_run = p.add_run(content)
                self._set_run_font(content_run)
                p.paragraph_format.line_spacing = 1.5

    def _add_transcript_image(self, doc: Document, item: Dict[str, Any]):
        """Add image from transcript item"""
        try:
            img_path = Path(item.get("path", ""))
            if not img_path.exists():
                logger.warning(f"Image not found: {img_path}")
                return
            
            # Add image
            doc.add_picture(str(img_path), width=Inches(5))
            
            # Center the image
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Add caption in gray
            timestamp = self._format_timestamp(item.get("timestamp", 0))
            caption = item.get("caption", "")
            img_type = item.get("image_type", "general")
            type_label = self._get_type_label(img_type)
            
            caption_text = f"{type_label}[{timestamp}]"
            if caption:
                caption_text += f" {caption[:150]}"
            
            p = doc.add_paragraph()
            run = p.add_run(caption_text)
            self._set_run_font(run, size=FONT_SIZE_CAPTION, color=COLOR_GRAY, italic=True)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            doc.add_paragraph()  # Space after image
            
        except Exception as e:
            logger.error(f"Failed to add image: {e}")

    def _add_section(self, doc: Document, section: Dict[str, Any]):
        """Add a content section (legacy support)"""
        # Section heading with timestamp
        heading_text = f"{section['title']} [{self._format_timestamp(section['start_time'])}]"
        doc.add_heading(heading_text, level=3)

        # Add images
        for image in section.get("images", []):
            self._add_image(doc, image, section['start_time'])

        # Add content
        if section.get("content"):
            p = doc.add_paragraph(section["content"])
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = Inches(0.3)

        # Add key points
        if section.get("key_points"):
            doc.add_paragraph("要点:")
            for point in section["key_points"]:
                p = doc.add_paragraph(point, style='List Bullet')

        doc.add_paragraph()

    def _add_image(self, doc: Document, image: Dict[str, Any], section_time: float):
        """Add image with caption"""
        try:
            img_path = Path(image["path"])
            if not img_path.exists():
                logger.warning(f"Image not found: {img_path}")
                return

            # Add image
            doc.add_picture(str(img_path), width=Inches(5))

            # Add caption with type label
            img_type = image.get("type", "general")
            type_label = self._get_type_label(img_type)

            caption_text = f"{type_label} {self._format_timestamp(image['timestamp'])}"
            if image.get("caption"):
                caption_text += f" - {image['caption'][:100]}"

            p = doc.paragraphs[-1]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            caption = p.add_run(f"\n{caption_text}")
            self._set_run_font(caption, size=FONT_SIZE_CAPTION, color=COLOR_GRAY, italic=True)

            doc.add_paragraph()

        except Exception as e:
            logger.error(f"Failed to add image: {e}")

    def _add_statistics(self, doc: Document, statistics: Dict[str, Any]):
        """Add statistics section"""
        heading = doc.add_heading("统计信息", level=2)
        for run in heading.runs:
            self._set_run_font(run, size=FONT_SIZE_HEADING, bold=True)

        stats = [
            f"视频时长: {self._format_duration(statistics.get('total_duration', 0))}",
            f"转录片段: {statistics.get('total_segments', 0)}",
            f"截取帧数: {statistics.get('total_frames', 0)}",
            f"成功分析: {statistics.get('successful_analyses', 0)}",
        ]

        for stat in stats:
            p = doc.add_paragraph(stat, style='List Bullet')
            for run in p.runs:
                self._set_run_font(run)

    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename for Windows"""
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, '_')
        return filename[:100] or "untitled"

    def _format_duration(self, seconds: float) -> str:
        """Format duration as HH:MM:SS"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)

        if hours > 0:
            return f"{hours}h {minutes}m {secs}s"
        elif minutes > 0:
            return f"{minutes}m {secs}s"
        else:
            return f"{secs}s"

    def _format_timestamp(self, seconds: float) -> str:
        """Format timestamp as HH:MM:SS"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)

        return f"{hours:02d}:{minutes:02d}:{secs:02d}"

    def _get_type_label(self, img_type: str) -> str:
        """Get type label for image"""
        labels = {
            "formula": "【公式】",
            "code": "【代码】",
            "chart": "【图表】",
            "text": "【文字】",
            "general": "",
        }
        return labels.get(img_type, "")


def generate_docx(data: Dict[str, Any], output_path: str) -> Path:
    """
    Convenience function to generate Word document

    Args:
        data: Structured note data
        output_path: Output file path

    Returns:
        Path to generated document
    """
    generator = DocxGenerator()
    return generator.generate(data, Path(output_path).stem)
